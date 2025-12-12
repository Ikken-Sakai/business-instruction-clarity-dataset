#!/usr/bin/env python3
"""
中間発表 口頭ノート（Word文書）生成スクリプト
外国人労働者のための日本語ビジネス指示文の曖昧性判定システム
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """水平線を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)

def main():
    """メイン処理"""
    doc = Document()
    
    # フォント設定用のスタイル
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(11)
    
    # ========================================
    # タイトル
    # ========================================
    title = doc.add_heading('中間発表 口頭ノート', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('外国人労働者のための日本語ビジネス指示文の曖昧性判定システム')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # ========================================
    # 発表時間配分
    # ========================================
    doc.add_heading('発表時間配分（合計: 約3分）', level=1)
    
    time_table = doc.add_table(rows=7, cols=3)
    time_table.style = 'Table Grid'
    
    headers = ['スライド', '内容', '時間']
    for i, header in enumerate(headers):
        cell = time_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    time_data = [
        ('1', 'タイトル', '15秒'),
        ('2', '背景・目的', '45秒'),
        ('3', '方法', '45秒'),
        ('4', '結果（進捗報告）', '40秒'),
        ('5', 'まとめ・考察', '35秒'),
        ('合計', '', '約3分'),
    ]
    
    for row_idx, (slide, content, time) in enumerate(time_data, 1):
        time_table.rows[row_idx].cells[0].text = slide
        time_table.rows[row_idx].cells[1].text = content
        time_table.rows[row_idx].cells[2].text = time
        if slide == '合計':
            for cell in time_table.rows[row_idx].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    add_horizontal_line(doc)
    
    # ========================================
    # スライド1: タイトル
    # ========================================
    doc.add_heading('スライド1: タイトル（約15秒）', level=1)
    
    doc.add_heading('発表原稿', level=2)
    script1 = doc.add_paragraph()
    script1.paragraph_format.left_indent = Inches(0.3)
    run = script1.add_run(
        '「これより、『外国人労働者のための日本語ビジネス指示文の曖昧性判定システム』について発表します。」'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    add_horizontal_line(doc)
    
    # ========================================
    # スライド2: 背景・目的
    # ========================================
    doc.add_heading('スライド2: 背景・目的（約45秒）', level=1)
    
    doc.add_heading('発表原稿', level=2)
    
    scripts2 = [
        '「日本の職場では、『例の件、早めに対応しといて』や『いい感じにまとめておいて』といった曖昧な指示が日常的に使われています。',
        '',
        '安部（2018）の研究では、留学生が日本の職場で感じる違和感として『指示内容の曖昧さ』や『暗黙の了解』が上位に挙がっていることが報告されています。',
        '',
        'また、宗（2021）の博士論文では、日本人が『ウチ』の関係に対して間接的・曖昧な表現を用いる傾向があり、これが外国人労働者とのコミュニケーションギャップの原因になっていることが指摘されています。',
        '',
        'そこで本研究では、BERTを用いた二値分類モデルによって、指示文が『明確』か『曖昧』かを自動判定し、曖昧な指示をリアルタイムで検出・改善支援するシステムの構築を目指します。」',
    ]
    
    for script in scripts2:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(script)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_heading('補足情報（質疑応答用）', level=2)
    supp2 = doc.add_paragraph()
    supp2.paragraph_format.left_indent = Inches(0.3)
    run = supp2.add_run('・安部論文の詳細: 地球社会統合科学研究 第8号 pp.1-12\n・宗論文の詳細: 日本大学博士論文、ウチソト関係の文化的背景を詳述')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_horizontal_line(doc)
    
    # ========================================
    # スライド3: 方法
    # ========================================
    doc.add_heading('スライド3: 方法（約45秒）', level=1)
    
    doc.add_heading('発表原稿', level=2)
    
    scripts3 = [
        '「手法としては、東北大学の日本語BERTモデルを使用した二値分類を行います。入力はビジネス指示文テキスト、出力は『明確（Label 0）』または『曖昧（Label 1）』です。',
        '',
        '学習データセットは自作しました。総データ数2,000件で、明確と曖昧が50:50の均等な分布です。',
        '',
        'データには、10業種のシチュエーション、丁寧語・タメ口・命令形の3種類の口調を含め、実際の職場で使われる多様な表現を網羅しました。',
        '',
        'なぜデータセットを自作したかというと、『ビジネス指示文の曖昧性』に特化した既存データセットが存在しなかったためです。東中ら（2016）の対話破綻コーパス構築手法を参考に、明確なアノテーション基準を策定しました。」',
    ]
    
    for script in scripts3:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(script)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_heading('補足情報（質疑応答用）', level=2)
    supp3 = doc.add_paragraph()
    supp3.paragraph_format.left_indent = Inches(0.3)
    run = supp3.add_run(
        '・BERT選定理由: 文脈を考慮した埋め込み表現が可能、日本語対応モデルが利用可能\n'
        '・データセット形式: JSONL（text, label, reason）\n'
        '・従来手法との違い: 奥村・田中（1989）のルールベースと異なり、文脈化埋め込みでニュアンスまで捉える'
    )
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_horizontal_line(doc)
    
    # ========================================
    # スライド4: 結果
    # ========================================
    doc.add_heading('スライド4: 結果（約40秒）', level=1)
    
    doc.add_heading('発表原稿', level=2)
    
    scripts4 = [
        '「現在の進捗ですが、データセット構築は完了しています。2,000件のデータを作成し、train/val/testに8:1:1で分割しました。',
        '',
        'データの品質チェックも行い、ラベルバランスは完全に50:50、文字数は平均24.2文字で、実際のビジネス指示に近い長さになっています。',
        '',
        '一方、BERTモデルの学習とデモアプリ開発は未着手です。',
        '',
        'まだモデル学習に着手していない理由は、まずデータセットの品質を十分に担保することを優先したためです。」',
    ]
    
    for script in scripts4:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(script)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 128)
    
    doc.add_heading('補足情報（質疑応答用）', level=2)
    supp4 = doc.add_paragraph()
    supp4.paragraph_format.left_indent = Inches(0.3)
    run = supp4.add_run(
        '・品質チェックスクリプト: check_dataset.py\n'
        '・ファイル構成: dataset.jsonl(2000件), train.jsonl(1600件), val.jsonl(200件), test.jsonl(200件)'
    )
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_horizontal_line(doc)
    
    # ========================================
    # スライド5: まとめ・考察
    # ========================================
    doc.add_heading('スライド5: まとめ・考察（約35秒）', level=1)
    
    doc.add_heading('発表原稿', level=2)
    
    scripts5 = [
        '「まとめとして、現状は高品質なデータセットの構築が完了し、モデル学習の準備が整った段階です。',
        '',
        '予想される結果として、ラベルバランスが均等で判定基準が明確なため、F1スコア85%以上を目標としています。',
        '',
        '本番までの計画として、12月中旬までにBERT学習とパラメータ調整、12月下旬にテストデータでの評価と結果分析、1月上旬にGradioでのデモアプリ開発と発表準備を行います。',
        '',
        '以上で発表を終わります。ご清聴ありがとうございました。」',
    ]
    
    for script in scripts5:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(script)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 128)
    
    add_horizontal_line(doc)
    doc.add_page_break()
    
    # ========================================
    # 質疑応答用メモ
    # ========================================
    doc.add_heading('質疑応答用メモ', level=1)
    
    qa_list = [
        ('Q1: なぜ外国人労働者に特化したのか？',
         'A: 安部（2018）の研究で、留学生が職場で感じる違和感の上位に「指示の曖昧さ」があることが示されており、実際のニーズがあるため。'),
        ('Q2: データセットは実際の職場データではないのか？',
         'A: 自動生成データです。実際の職場データは個人情報や機密情報を含むため収集が困難でした。今後は実証実験で実データとの乖離を検証予定です。'),
        ('Q3: 日本人上司の指導にも使えるか？',
         'A: はい。曖昧な指示を出した際にリアルタイムでフィードバックすることで、コミュニケーション改善ツールとしても活用可能です。'),
        ('Q4: 従来の曖昧性解消手法との違いは？',
         'A: 奥村・田中（1989）のルールベース手法と異なり、BERTは文脈化埋め込みにより「なる早」「いい感じ」といったニュアンスまで捉えられます。'),
        ('Q5: 2,000件のデータで十分か？',
         'A: BERTはファインチューニング用途では比較的少量のデータでも効果を発揮します。まずはこの規模で学習し、必要に応じて拡張予定です。'),
        ('Q6: 「明確/曖昧」の二値分類で十分か？',
         'A: 今回は二値分類ですが、将来的には曖昧性の種類別分類（What欠如/When欠如/How欠如）への拡張も検討しています。'),
    ]
    
    for q, a in qa_list:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(q)
        q_run.font.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(0, 51, 102)
        
        a_para = doc.add_paragraph()
        a_para.paragraph_format.left_indent = Inches(0.3)
        a_run = a_para.add_run(a)
        a_run.font.size = Pt(11)
        
        doc.add_paragraph()
    
    add_horizontal_line(doc)
    
    # ========================================
    # 参考文献
    # ========================================
    doc.add_heading('参考文献（フル情報）', level=1)
    
    references = [
        '1. 安部陽子 (2018). 「日本での就労時に留学生が持つ違和感の調査報告：日本人学生との対照分析を通して」, 地球社会統合科学研究, 第8号, pp.1-12.',
        '2. 宗甜甜 (2021). 「ビジネス場面における日本語の『断り』に関する研究」, 博士論文（日本大学）.',
        '3. 東中竜一郎 他 (2016). 「テキストチャットを用いた雑談対話コーパスの構築と対話破綻の分析」, 自然言語処理, Vol.23, No.1.',
        '4. 奥村学, 田中穂積 (1989). 「自然言語解析における意味的曖昧性を増進的に解消する計算モデル」, 人工知能学会誌, Vol.4, No.6.',
        '5. Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding." arXiv preprint arXiv:1810.04805.',
        '6. 李婷 (2023). 「ビジネスメールにおけるメタ言語表現から読み取れる待遇意識」.',
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
    
    # ========================================
    # 保存
    # ========================================
    output_path = os.path.join(os.path.dirname(__file__), "中間発表ノート_外国人労働者向け曖昧性判定システム.docx")
    doc.save(output_path)
    print(f"✅ Word発表ノートを生成しました: {output_path}")

if __name__ == "__main__":
    main()




















